# src/export.py
from __future__ import annotations

from typing import List, Optional, Any, Dict
import json
from io import BytesIO

from docx import Document

from .ocr_engine import OcrLine


def _safe_lines(lines: Optional[List[OcrLine]]) -> List[OcrLine]:
    if not lines:
        return []
    # lọc None + đảm bảo có thuộc tính text
    out = []
    for ln in lines:
        if ln is None:
            continue
        out.append(ln)
    return out


def lines_to_text(lines: Optional[List[OcrLine]]) -> str:
    """
    Safe: nếu lines None/[] -> trả ""
    """
    lines = _safe_lines(lines)
    return "\n".join([(getattr(ln, "text", "") or "").strip() for ln in lines]).strip()


def lines_to_json(lines: Optional[List[OcrLine]], ensure_ascii: bool = False) -> str:
    """
    Safe JSON: gồm text + bbox/quad nếu có
    """
    lines = _safe_lines(lines)
    payload: List[Dict[str, Any]] = []
    for ln in lines:
        item: Dict[str, Any] = {"text": (getattr(ln, "text", "") or "")}
        if hasattr(ln, "bbox") and ln.bbox is not None:
            item["bbox"] = ln.bbox
        if hasattr(ln, "quad") and ln.quad is not None:
            item["quad"] = ln.quad
        payload.append(item)
    return json.dumps(payload, ensure_ascii=ensure_ascii, indent=2)


def lines_to_docx_bytes(lines: Optional[List[OcrLine]]) -> bytes:
    """
    Safe DOCX: nếu rỗng vẫn tạo file docx hợp lệ.
    """
    lines = _safe_lines(lines)
    doc = Document()
    if not lines:
        doc.add_paragraph("")  # doc rỗng vẫn hợp lệ
    else:
        for ln in lines:
            doc.add_paragraph((getattr(ln, "text", "") or "").strip())
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
